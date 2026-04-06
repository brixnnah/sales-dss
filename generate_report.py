"""
Generate TECHNICAL_REPORT.docx from report content.
Run: python generate_report.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────
def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x11, 0x62, 0x7d)
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    return p

def para(text='', bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    r = p.add_run(text)
    set_font(r)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, size=10, color=(0xFF, 0xFF, 0xFF))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A56DB')
        tcPr.append(shd)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=10)
            if ri % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EFF6FF')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('Technical Report')
set_font(r, size=28, bold=True, color=(0x1a, 0x56, 0xdb))

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle_p.add_run('Sales Forecasting & Inventory Optimization\nDecision Support System')
set_font(r2, size=18, color=(0x11, 0x62, 0x7d))

doc.add_paragraph()

meta = [
    ('Project Title', 'Sales Forecasting & Inventory Optimization DSS'),
    ('System Type',   'Web-Based Decision Support System'),
    ('Stack',         'Python · Flask · Prophet · Pandas · Chart.js'),
    ('Date',          'April 2026'),
    ('Team Members',  '[Student 1 Name] · [Student 2 Name] · [Student 3 Name]'),
]
add_table(['Field', 'Detail'], meta, col_widths=[1.5, 4.0])
doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═════════════════════════════════════════════════════════════
heading('Table of Contents', 1)
toc_items = [
    '1.  Introduction',
    '2.  Literature Background',
    '3.  Methodology',
    '4.  System Architecture',
    '5.  Decision Model Explanation',
    '6.  Results and Evaluation',
    '7.  Challenges and Limitations',
    '8.  Conclusion and Recommendations',
    '9.  References',
]
for item in toc_items:
    bullet(item)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═════════════════════════════════════════════════════════════
heading('1. Introduction', 1)
heading('1.1 Background and Motivation', 2)
para(
    'Retail and supermarket operations generate large volumes of transactional data daily. '
    'Despite this data richness, many managers continue to rely on intuition or static spreadsheets '
    'when making inventory replenishment decisions. Poor decisions directly translate to lost revenue '
    'through stockouts or wasted capital through overstock. A structured, data-driven approach to these '
    'decisions is both feasible and necessary.'
)
para(
    'This project addresses this gap by designing and implementing a functional Decision Support System '
    '(DSS) tailored to sales forecasting and inventory optimization. The system ingests historical sales '
    'transaction data, applies statistical forecasting to project future demand, and generates rule-based '
    'inventory recommendations that a manager can act on immediately.'
)

heading('1.2 Problem Statement', 2)
para(
    'Given historical sales data across multiple product categories, how can a retail manager determine '
    'how much of each category to reorder, and when, to minimize stockout risk while avoiding excess '
    'inventory? This is a recurring, semi-structured decision that benefits from system support. Without '
    'a DSS, managers must manually aggregate data, estimate trends, and apply arbitrary buffer rules — '
    'all of which are error-prone and inconsistent across periods.',
    italic=True
)

heading('1.3 Decision-Makers', 2)
for dm in [
    'Inventory Managers — responsible for purchase order decisions',
    'Store Operations Managers — accountable for shelf availability and waste',
    'Category Managers — focused on performance of individual product groups',
]:
    bullet(dm)
doc.add_paragraph()

heading('1.4 System Objectives', 2)
for obj in [
    'Transform raw transactional data into a clean, analysis-ready dataset.',
    'Provide an interactive dashboard visualizing key performance indicators (KPIs).',
    'Perform exploratory demand analysis across categories and time periods.',
    'Generate short-term (7–90 day) demand forecasts per product category.',
    'Compute actionable inventory recommendations (reorder point, safety stock, target stock).',
    'Secure system access via user authentication for authorized users only.',
]:
    bullet(obj)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 2. LITERATURE BACKGROUND
# ═════════════════════════════════════════════════════════════
heading('2. Literature Background', 1)
heading('2.1 Decision Support Systems: Core Concepts', 2)
para(
    'A Decision Support System (DSS) is an interactive information system that assists decision-makers '
    'in using data and models to solve semi-structured problems (Keen & Morton, 1978). Sprague and '
    'Carlson (1982) identified three fundamental DSS components:'
)
for c in [
    'Database Management Subsystem — stores and retrieves relevant data',
    'Model-Based Management Subsystem — applies analytical or predictive logic',
    'User Interface Subsystem — presents outputs and accepts inputs from decision-makers',
]:
    bullet(c)
para(
    '\nThis project maps directly to this framework: the cleaned CSV dataset serves as the data layer, '
    'Prophet forecasting and the inventory formulas form the model layer, and the Flask web application '
    'provides the interactive interface.'
)

heading('2.2 DSS in Retail and Inventory Management', 2)
para(
    'Retail DSS applications have been studied extensively (Coyle et al., 2013). Common models applied include:'
)
for m in [
    'Time Series Forecasting — used to project future demand from historical patterns (Box & Jenkins, 1976; Taylor, 2018)',
    'Reorder Point (ROP) Models — determine when to reorder based on lead time and average demand',
    'Safety Stock Models — statistical buffers based on demand variability and service level targets (Silver et al., 2017)',
]:
    bullet(m)

heading('2.3 Time Series Forecasting with Prophet', 2)
para(
    'Prophet (Taylor & Letham, 2018), developed by Meta, is a decomposable time series model designed '
    'for business data. It decomposes time series as:'
)
para('y(t) = g(t) + s(t) + h(t) + ε(t)', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(
    'Where g(t) is the trend, s(t) is the seasonality component, h(t) represents holiday effects, and '
    'ε(t) is a normally distributed error term. Prophet handles missing data, multiple seasonality '
    'patterns, and provides interpretable uncertainty intervals — making it well-suited for inventory planning.'
)

heading('2.4 Inventory Theory', 2)
para('The following classical formulas underpin the recommendation engine:')
for f in [
    'Reorder Point (ROP) = Average Daily Demand × Lead Time',
    'Safety Stock (SS) = z × σ_demand × √(Lead Time)',
    'Target Stock (TS) = ROP + SS',
    'Where z = 1.65 for a 95% service level',
]:
    bullet(f)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 3. METHODOLOGY
# ═════════════════════════════════════════════════════════════
heading('3. Methodology', 1)
heading('3.1 Research Approach', 2)
para(
    'This project follows a Design Science Research (DSR) methodology (Hevner et al., 2004), which '
    'emphasizes the creation and evaluation of a functional IT artifact. The five-step process was:'
)
for s in [
    'Problem identification — defined the inventory decision problem',
    'Data acquisition and cleaning — collected and preprocessed real transaction data',
    'Model development — implemented forecasting and inventory logic',
    'System construction — integrated models into a usable web application',
    'Evaluation — validated system outputs and tested all routes',
]:
    bullet(s)

heading('3.2 Dataset Description', 2)
add_table(
    ['Attribute', 'Value'],
    [
        ('Raw dataset rows', '286,392'),
        ('Time period', 'October 2020 – September 2021'),
        ('Key fields', 'order_date, category, qty_ordered, price, status, sku'),
        ('Source format', 'CSV'),
    ],
    col_widths=[2.0, 3.5]
)

heading('3.3 Data Cleaning Pipeline', 2)
para('The cleaning pipeline in cleaning.py performs 8 sequential steps:')
steps = [
    ('Status Filtering', 'Retain only received and complete orders; exclude cancelled or refunded records.'),
    ('PII Column Removal', 'Drop personally identifiable fields (name, email, SSN, phone, address) and irrelevant noise columns.'),
    ('Missing Value Removal', 'Drop rows missing any of: qty_ordered, price, order_date, category.'),
    ('Invalid Value Removal', 'Remove records with qty_ordered ≤ 0 or price ≤ 0 (data entry errors).'),
    ('Date Conversion', 'Parse order_date strings to Python datetime format for time-based operations.'),
    ('Category-Day Aggregation', 'Aggregate to category × date level, summing quantities and averaging prices.'),
    ('Outlier Capping', 'Cap daily quantities above the 99th percentile per category to suppress spikes.'),
    ('Rare Category Removal', 'Remove categories appearing on fewer than 5 distinct days.'),
]
add_table(
    ['Step', 'Name', 'Action'],
    [(str(i+1), n, d) for i, (n, d) in enumerate(steps)],
    col_widths=[0.4, 1.6, 3.5]
)

heading('3.4 Cleaning Results', 2)
add_table(
    ['Metric', 'Value'],
    [
        ('Rows after cleaning', '4,683'),
        ('Active categories', '15'),
        ('Date range (cleaned)', '2020-10-01 to 2021-09-29'),
        ('Row reduction', '~98.4% (aggregation from order-line to category-day)'),
    ],
    col_widths=[2.2, 3.3]
)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 4. SYSTEM ARCHITECTURE
# ═════════════════════════════════════════════════════════════
heading('4. System Architecture', 1)
heading('4.1 Overview', 2)
para(
    'The system follows a three-tier Model-View-Controller (MVC) architecture deployed as a local '
    'Flask web application:'
)
for tier in [
    'Presentation Layer — HTML/CSS templates with Chart.js visualizations (index.html, eda.html, forecast.html, login.html)',
    'Application Layer — Flask application (app.py) handling routes, authentication, data loading, and model execution',
    'Data Layer — Raw CSV (sales.csv) → cleaning pipeline (cleaning.py) → cleaned CSV (sales_CLEANED.csv)',
]:
    bullet(tier)

heading('4.2 Component Description', 2)
add_table(
    ['Component', 'File', 'Responsibility'],
    [
        ('Data Cleaning Pipeline', 'cleaning.py', 'Raw CSV → cleaned daily category sales'),
        ('Flask Application', 'app.py', 'Route handling, data loading, model execution'),
        ('Dashboard Template', 'Templates/index.html', 'KPI cards, charts, DSS recommendations table'),
        ('EDA Template', 'Templates/eda.html', 'Statistical profile, trend charts, insights, print export'),
        ('Forecast Template', 'Templates/forecast.html', 'Forecast output, inventory recommendation table'),
        ('Login Template', 'Templates/login.html', 'User authentication form'),
        ('Stylesheet', 'static/styles.css', 'Unified design system for all pages'),
        ('Tests', 'tests/test_app.py', 'Automated regression tests for core logic'),
    ],
    col_widths=[1.8, 1.8, 2.9]
)

heading('4.3 Application Routes', 2)
add_table(
    ['Route', 'Method', 'Auth Required', 'Purpose'],
    [
        ('/login',    'GET, POST', 'No',  'User authentication form and login processing'),
        ('/logout',   'GET',       'Yes', 'End session and redirect to login'),
        ('/',         'GET',       'Yes', 'Main KPI dashboard with DSS recommendations'),
        ('/eda',      'GET',       'Yes', 'Exploratory data analysis with category filter'),
        ('/forecast', 'GET',       'Yes', 'Prophet forecast and inventory recommendation'),
    ],
    col_widths=[1.0, 1.1, 1.2, 3.2]
)

heading('4.4 Security Design', 2)
para(
    'User authentication is implemented using Flask-Login with server-side session management. '
    'All three data routes are decorated with @login_required, ensuring unauthenticated requests '
    'are automatically redirected to /login before any data is served. Two roles are supported: '
    'Administrator and Manager.'
)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 5. DECISION MODEL EXPLANATION
# ═════════════════════════════════════════════════════════════
heading('5. Decision Model Explanation', 1)
heading('5.1 Model Overview', 2)
para('The DSS integrates two decision models:')
for m in [
    'Predictive Model — Prophet time series forecasting for future demand estimation',
    'Rule-Based Inventory Model — classical reorder point and safety stock formulas for replenishment recommendations',
]:
    bullet(m)
para(
    'These two models are integrated so that forecast output directly feeds the inventory '
    'recommendation computation.'
)

heading('5.2 Prophet Forecasting Model', 2)
para('Input: A category-filtered time series of daily qty_ordered values.')
para('Configuration used:')
for c in [
    'Weekly seasonality: enabled (demand varies meaningfully by day of week)',
    'Yearly seasonality: disabled (only 12 months of data — insufficient for annual pattern detection)',
    'Daily seasonality: disabled (pre-aggregated daily inputs)',
    'Confidence interval width: 90%',
    'Minimum data requirement: 7 days of category history',
]:
    bullet(c)
para('Output per forecast record:')
add_table(
    ['Field', 'Description'],
    [
        ('date',        'Forecast date'),
        ('forecast',    'Point estimate of expected daily demand quantity (yhat)'),
        ('lower_bound', '90% lower confidence bound (yhat_lower, clamped to 0)'),
        ('upper_bound', '90% upper confidence bound (yhat_upper)'),
    ],
    col_widths=[1.5, 4.0]
)

heading('5.3 Inventory Recommendation Engine', 2)
para('The inventory engine computes three metrics per selected category:')
add_table(
    ['Metric', 'Formula', 'Default Inputs'],
    [
        ('Average Daily Demand (d̄)', 'Mean of forecast quantities (clamped ≥ 0)', 'Forecast records'),
        ('Reorder Point (ROP)',       'ROP = d̄ × L',                               'L = 14 days'),
        ('Safety Stock (SS)',         'SS = z × σ_d × √L',                          'z = 1.65 (95% service level)'),
        ('Target Stock (TS)',         'TS = ROP + SS',                               '—'),
    ],
    col_widths=[1.8, 2.2, 1.5]
)

heading('5.4 Dashboard Recommendation Labels', 2)
para('The home dashboard applies a buffer-based rule for the top-N category overview:')
para('Recommended Reorder = Monthly Average × (1 + 0.10 × Buffer Months)', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table(
    ['Condition', 'Status Label'],
    [
        ('Recommended Reorder > Monthly Average × 1.5', '🔴 Low Stock'),
        ('Recommended Reorder > Monthly Average × 1.2', '🟡 Monitor'),
        ('Otherwise',                                    '🟢 Optimal'),
    ],
    col_widths=[3.5, 2.0]
)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 6. RESULTS AND EVALUATION
# ═════════════════════════════════════════════════════════════
heading('6. Results and Evaluation', 1)
heading('6.1 Dataset Summary Statistics', 2)
add_table(
    ['Metric', 'Value'],
    [
        ('Total observations (cleaned)', '4,683 rows'),
        ('Active product categories',    '15'),
        ('Analysis period',              'Oct 2020 – Sep 2021 (12 months)'),
        ('Total gross revenue',          '$192,119,084'),
        ('Total units sold',             '416,774'),
        ('Mean unit price',              '$334.86'),
    ],
    col_widths=[2.5, 3.0]
)

heading('6.2 Top Categories by Sales Volume', 2)
add_table(
    ['Rank', 'Category', 'Total Units'],
    [
        ('1', 'Superstore',        '61,139'),
        ('2', 'Mobiles & Tablets', '55,202'),
        ('3', "Men's Fashion",     '51,518'),
        ('4', "Women's Fashion",   '42,123'),
        ('5', 'Appliances',        '40,860'),
    ],
    col_widths=[0.6, 2.5, 2.4]
)

heading('6.3 Top Categories by Revenue', 2)
add_table(
    ['Rank', 'Category', 'Total Revenue'],
    [
        ('1', 'Mobiles & Tablets', '$86,357,322'),
        ('2', 'Appliances',        '$41,743,400'),
        ('3', 'Entertainment',     '$29,119,839'),
        ('4', 'Computing',         '$7,840,945'),
        ('5', 'Others',            '$6,319,471'),
    ],
    col_widths=[0.6, 2.5, 2.4]
)
para(
    'Key insight: Mobiles & Tablets ranks 2nd in volume but 1st in revenue by a large margin, '
    'indicating high average unit prices. This category should be prioritized for safety stock '
    'maintenance due to its disproportionate revenue contribution.',
    italic=True
)

heading('6.4 Sample Forecast Output — Appliances Category', 2)
add_table(
    ['Metric', 'Value'],
    [
        ('Average Monthly Demand', '3,405 units'),
        ('Lead Time',              '14 days'),
        ('Reorder Point',          '1,598 units'),
        ('Safety Stock',           '1,474 units'),
        ('Target Stock',           '3,072 units'),
    ],
    col_widths=[2.5, 3.0]
)
para(
    'Interpretation: A manager should initiate a reorder when current Appliances stock falls to '
    '1,598 units, maintaining a 1,474-unit buffer to absorb demand variability during the 14-day '
    'replenishment lead time.'
)

heading('6.5 System Feature Evaluation', 2)
add_table(
    ['Feature', 'Status', 'Notes'],
    [
        ('Data loading and cleaning',  '✅ Operational', '286,392 → 4,683 rows after cleaning'),
        ('Dashboard KPI display',      '✅ Operational', 'Revenue, profit, transactions, avg order value'),
        ('Date range filtering',       '✅ Operational', 'Applied across all three main routes'),
        ('Category filtering (EDA)',   '✅ Operational', '15 categories available via dropdown'),
        ('Prophet forecasting',        '✅ Operational', '7–90 day horizon, per category'),
        ('Inventory recommendation',   '✅ Operational', 'ROP, SS, target stock computed live'),
        ('User authentication',        '✅ Operational', 'Login/logout with role display in sidebar'),
        ('Automated tests',            '✅ Passing',     '3/3 tests pass'),
        ('Print/export (EDA)',         '✅ Operational', 'Print-friendly layout with report header'),
    ],
    col_widths=[2.0, 1.2, 2.3]
)

heading('6.6 Automated Test Results', 2)
add_table(
    ['Test', 'Description', 'Result'],
    [
        ('test_compute_daily_metrics_uses_category', 'Verifies metrics function returns category-level breakdown', '✅ PASS'),
        ('test_build_dss_recommendations',           'Verifies reorder recommendations computed correctly',       '✅ PASS'),
        ('test_load_data_has_order_date',            'Verifies data loader parses dates correctly',               '✅ PASS'),
    ],
    col_widths=[2.4, 2.6, 0.7]
)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 7. CHALLENGES AND LIMITATIONS
# ═════════════════════════════════════════════════════════════
heading('7. Challenges and Limitations', 1)
heading('7.1 Challenges Encountered', 2)
challenges = [
    ('Data Quality', 'The raw dataset contained significant noise: cancelled orders, null values, negative quantities, and low-frequency categories. The cleaning pipeline required multiple iterations to settle on category-day aggregation as the right granularity for forecasting.'),
    ('Forecast Horizon vs. Data Volume', 'With only 12 months of historical data, yearly seasonality could not be reliably modeled. Prophet\'s yearly seasonality was explicitly disabled to prevent the model from fitting spurious annual patterns.'),
    ('Negative Confidence Intervals', 'Prophet\'s uncertainty intervals can extend below zero for low-demand categories. The inventory engine clamps lower bounds at zero, which is statistically appropriate but requires explanation to non-technical reviewers.'),
    ('Static User Store', 'For this prototype, user credentials are stored in plain text in memory. A production deployment would require a database-backed hashed credential store.'),
    ('Optional Plotly Warning', 'Prophet optionally imports Plotly for interactive plots. Since Plotly is not installed, a non-critical warning appears at startup, which has no effect on system functionality.'),
]
for title, desc in challenges:
    p = doc.add_paragraph()
    r1 = p.add_run(title + ': ')
    set_font(r1, bold=True)
    r2 = p.add_run(desc)
    set_font(r2)

heading('7.2 System Limitations', 2)
add_table(
    ['Limitation', 'Impact', 'Mitigation'],
    [
        ('Category-level forecasting only',  'Cannot support SKU-level reorder decisions',          'Documented as future extension'),
        ('Fixed lead time (14 days)',         'Does not reflect supplier variability',               'Parameter exposed in backend for developer extension'),
        ('95% service level fixed in UI',    'Manager cannot adjust risk tolerance interactively',   'Logic is parameterized in backend functions'),
        ('Single-year data',                 'Long-term seasonal patterns undetectable',             'Yearly seasonality disabled in Prophet config'),
        ('No inventory balance input',       'Recommendations assume no current stock context',      'System is advisory; manager supplies current stock context'),
    ],
    col_widths=[1.8, 1.8, 2.0]
)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 8. CONCLUSION AND RECOMMENDATIONS
# ═════════════════════════════════════════════════════════════
heading('8. Conclusion and Recommendations', 1)
heading('8.1 Conclusion', 2)
para(
    'This project successfully delivered a functional Decision Support System for retail sales '
    'forecasting and inventory optimization. The system integrates all required DSS components: a '
    'data pipeline for preprocessing, a predictive model (Prophet) for demand forecasting, a '
    'rule-based inventory engine for replenishment decisions, and an interactive Flask web application '
    'for decision-maker engagement.'
)
para(
    'The system addresses a real and recurring semi-structured decision problem — inventory '
    'replenishment — where data-driven forecasting and business rules meaningfully improve decision '
    'quality over unaided human judgment. All three application pages (Dashboard, EDA, Forecast) are '
    'functional and provide complementary views: retrospective performance, statistical demand '
    'profiling, and forward-looking decision support.'
)
para(
    'The automated test suite confirms that core decision logic functions correctly, user authentication '
    'ensures appropriate access control, and the print/export capability of the EDA page makes the '
    'system directly usable for operational reporting.'
)

heading('8.2 Recommendations for Extension', 2)
para('Short-term:', bold=True)
for r in [
    'Introduce a database backend (SQLite or PostgreSQL) to replace the in-memory user store and enable audit logging of decisions.',
    'Allow the manager to input current stock on hand on the Forecast page so the system can calculate days-of-stock-remaining.',
    'Make lead time and service level user-editable on the forecast form to enable live scenario comparison.',
]:
    bullet(r)

para('Medium-term:', bold=True)
for r in [
    'Extend forecasting to SKU level by clustering SKUs into demand groups for more granular replenishment decisions.',
    'Incorporate external signals (promotions, holidays) as Prophet regressors to improve forecast accuracy around peaks.',
    'Add a scenario simulator where managers can compare inventory costs under different reorder strategies.',
]:
    bullet(r)

para('Long-term:', bold=True)
for r in [
    'Connect to a live ERP or POS data feed for real-time decision support rather than batch CSV processing.',
    'Implement an Economic Order Quantity (EOQ) model to also optimize order size, not just timing.',
    'Add email or notification triggers when a category is forecast to drop below its reorder point within the lead time window.',
]:
    bullet(r)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════
# 9. REFERENCES
# ═════════════════════════════════════════════════════════════
heading('9. References', 1)
refs = [
    'Box, G. E. P., & Jenkins, G. M. (1976). Time Series Analysis: Forecasting and Control. Holden-Day.',
    'Coyle, J. J., Langley, C. J., Novack, R. A., & Gibson, B. J. (2013). Supply Chain Management: A Logistics Perspective (9th ed.). Cengage Learning.',
    'Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design science in information systems research. MIS Quarterly, 28(1), 75–105.',
    'Keen, P. G. W., & Morton, M. S. S. (1978). Decision Support Systems: An Organizational Perspective. Addison-Wesley.',
    'Silver, E. A., Pyke, D. F., & Thomas, D. J. (2017). Inventory and Production Management in Supply Chains (4th ed.). CRC Press.',
    'Sprague, R. H., & Carlson, E. D. (1982). Building Effective Decision Support Systems. Prentice-Hall.',
    'Taylor, S. J., & Letham, B. (2018). Forecasting at scale. The American Statistician, 72(1), 37–45.',
]
for ref in refs:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(ref)
    set_font(r, size=10)

# Footer note
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run('Report generated: April 2026  |  System: Sales Forecasting & Inventory Optimization DSS  |  Framework: Flask + Prophet')
set_font(fr, size=9, italic=True, color=(0x6B, 0x72, 0x80))

# ── Save ──────────────────────────────────────────────────────
doc.save('TECHNICAL_REPORT.docx')
print('✅  TECHNICAL_REPORT.docx saved successfully.')
